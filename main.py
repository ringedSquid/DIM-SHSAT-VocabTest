from docx import Document 
import os

vocab_path = "Words.txt"
template_path = "Tempalte.docx"
test_dir = "tests/"

if (os.path.exists("tests/") == False):
    os.mkdir("tests/")

def make_tests(d_path, v_path, t_path):
    data = []
    doc = Document(d_path)

    with open(v_path, "r") as file:
        data = file.read().split("\n")

    for i in range(0, len(data)-10, 10):
        print("Processing test " + str(i//10 + 1))

        for y in range(10):
            doc.paragraphs[y*2 + 4].text = data[i+y] + ": "
            doc.paragraphs[y*2 + 4].text += "_" * (70 - len(doc.paragraphs[y*2 + 4].text))

        doc.save(t_path + "Vocab Test " + str(i//10 + 1) + ".docx")

make_tests("Template.docx", vocab_path, test_dir)

    



